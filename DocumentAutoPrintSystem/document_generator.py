from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib import colors
from docx import Document as DocxDocument
from datetime import datetime
import os
from database import db_manager


class DocumentGenerator:
    """Класс для генерации документов"""

    def __init__(self, output_dir='output'):
        self.output_dir = output_dir
        # Создаем папки для выходных файлов
        os.makedirs(os.path.join(output_dir, 'protocols'), exist_ok=True)
        os.makedirs(os.path.join(output_dir, 'resolutions'), exist_ok=True)

    def generate_protocol(self, data):
        """Генерация протокола в формате PDF"""
        # Формируем имя файла
        filename = f"Протокол_{data.get('number', '001')}_{datetime.now().strftime('%Y%m%d')}.pdf"
        filepath = os.path.join(self.output_dir, 'protocols', filename)

        # Создаем PDF документ
        doc = SimpleDocTemplate(
            filepath,
            pagesize=A4,
            leftMargin=2 * cm,
            rightMargin=2 * cm,
            topMargin=2 * cm,
            bottomMargin=2 * cm
        )

        # Стили документа
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            alignment=TA_CENTER,
            spaceAfter=30
        )

        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=12,
            alignment=TA_LEFT,
            spaceAfter=12
        )

        # Содержимое документа
        story = []

        # Заголовок
        story.append(Paragraph("ПРОТОКОЛ", title_style))
        story.append(Paragraph(f"№ {data.get('number', '')} от {data.get('date', '')}", title_style))
        story.append(Spacer(1, 20))

        # Основное содержание
        story.append(Paragraph(f"<b>Тема:</b> {data.get('topic', '')}", normal_style))
        story.append(Paragraph(f"<b>Место проведения:</b> {data.get('location', '')}", normal_style))
        story.append(Paragraph(f"<b>Дата и время:</b> {data.get('datetime', '')}", normal_style))
        story.append(Spacer(1, 20))

        # Присутствующие
        if 'participants' in data:
            story.append(Paragraph("<b>Присутствовали:</b>", normal_style))
            for participant in data['participants']:
                story.append(Paragraph(f"  • {participant}", normal_style))

        story.append(Spacer(1, 20))

        # Повестка дня
        if 'agenda' in data:
            story.append(Paragraph("<b>Повестка дня:</b>", normal_style))
            for i, item in enumerate(data['agenda'], 1):
                story.append(Paragraph(f"  {i}. {item}", normal_style))

        story.append(Spacer(1, 20))

        # Решения
        if 'decisions' in data:
            story.append(Paragraph("<b>Принятые решения:</b>", normal_style))
            for i, decision in enumerate(data['decisions'], 1):
                story.append(Paragraph(f"  {i}. {decision}", normal_style))

        story.append(Spacer(1, 30))

        # Подписи
        signatures = [
            ["Председатель:", "_________________", data.get('chairman', '')],
            ["Секретарь:", "_________________", data.get('secretary', '')]
        ]

        signature_table = Table(signatures, colWidths=[4 * cm, 4 * cm, 8 * cm])
        signature_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))

        story.append(signature_table)

        # Создаем документ
        doc.build(story)

        # Сохраняем документ в базе данных
        document_data = {
            'document_type': 'protocol',
            'document_number': data.get('number', ''),
            'document_date': datetime.strptime(data.get('date', ''), '%d.%m.%Y') if data.get(
                'date') else datetime.now(),
            'title': f"Протокол №{data.get('number', '')}",
            'content': str(data),
            'file_path': filepath,
            'author': data.get('author', 'Система'),
            'status': 'final'
        }

        doc_id = db_manager.add_document(**document_data)

        return filepath, doc_id

    def generate_resolution(self, data):
        """Генерация постановления в формате DOCX"""
        # Формируем имя файла
        filename = f"Постановление_{data.get('number', '001')}_{datetime.now().strftime('%Y%m%d')}.docx"
        filepath = os.path.join(self.output_dir, 'resolutions', filename)

        # Создаем Word документ
        doc = DocxDocument()

        # Заголовок
        doc.add_heading('ПОСТАНОВЛЕНИЕ', 0)
        doc.add_paragraph(f"№ {data.get('number', '')} от {data.get('date', '')}")
        doc.add_paragraph()

        # Основное содержание
        doc.add_paragraph(f"Тема: {data.get('topic', '')}")
        doc.add_paragraph(f"Основание: {data.get('basis', '')}")
        doc.add_paragraph()

        # Преамбула
        doc.add_paragraph("Рассмотрев материалы дела, заслушав объяснения сторон,")
        doc.add_paragraph("ПОСТАНОВИЛ:")
        doc.add_paragraph()

        # Решения
        if 'resolutions' in data:
            for i, resolution in enumerate(data['resolutions'], 1):
                doc.add_paragraph(f"{i}. {resolution}", style='List Number')

        doc.add_paragraph()

        # Подписи
        doc.add_paragraph("Председатель комиссии:")
        doc.add_paragraph(f"_________________ / {data.get('chairman', '')} /")
        doc.add_paragraph()
        doc.add_paragraph("Члены комиссии:")

        if 'members' in data:
            for member in data['members']:
                doc.add_paragraph(f"_________________ / {member} /")

        # Сохраняем документ
        doc.save(filepath)

        # Сохраняем документ в базе данных
        document_data = {
            'document_type': 'resolution',
            'document_number': data.get('number', ''),
            'document_date': datetime.strptime(data.get('date', ''), '%d.%m.%Y') if data.get(
                'date') else datetime.now(),
            'title': f"Постановление №{data.get('number', '')}",
            'content': str(data),
            'file_path': filepath,
            'author': data.get('author', 'Система'),
            'status': 'final'
        }

        doc_id = db_manager.add_document(**document_data)

        return filepath, doc_id

    def print_document(self, filepath):
        """Печать документа"""
        try:
            # Для Windows
            if os.name == 'nt':
                os.startfile(filepath, "print")
            # Для Linux
            elif os.name == 'posix':
                os.system(f"lp {filepath}")
            else:
                print(f"Файл готов для печати: {filepath}")
            return True
        except Exception as e:
            print(f"Ошибка при печати: {e}")
            return False